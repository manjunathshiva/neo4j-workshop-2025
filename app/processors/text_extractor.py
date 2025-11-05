"""
Text extraction utilities for various document formats.
Supports PDF, TXT, DOCX, and MD file formats with error handling.
"""

import logging
import io
import re
from typing import Optional, Dict, Any, BinaryIO, Union
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import markdown

logger = logging.getLogger(__name__)

class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    TXT = "txt"
    DOCX = "docx"
    MD = "md"
    MARKDOWN = "markdown"

@dataclass
class ExtractionResult:
    """Result of text extraction from a document."""
    success: bool
    text: str
    metadata: Dict[str, Any]
    error_message: Optional[str] = None
    word_count: int = 0
    character_count: int = 0
    
    def __post_init__(self):
        """Calculate text statistics after initialization."""
        if self.success and self.text:
            self.word_count = len(self.text.split())
            self.character_count = len(self.text)

class TextExtractor:
    """
    Text extraction utility for various document formats.
    Handles PDF, TXT, DOCX, and MD files with comprehensive error handling.
    """
    
    def __init__(self, max_file_size_mb: int = 10):
        """
        Initialize text extractor.
        
        Args:
            max_file_size_mb: Maximum file size in MB to process
        """
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.supported_types = {
            '.pdf': DocumentType.PDF,
            '.txt': DocumentType.TXT,
            '.docx': DocumentType.DOCX,
            '.md': DocumentType.MD,
            '.markdown': DocumentType.MARKDOWN
        }
    
    def extract_from_file(self, file_path: Union[str, Path]) -> ExtractionResult:
        """
        Extract text from a file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        try:
            file_path = Path(file_path)
            
            # Check if file exists
            if not file_path.exists():
                return ExtractionResult(
                    success=False,
                    text="",
                    metadata={},
                    error_message=f"File not found: {file_path}"
                )
            
            # Check file size
            file_size = file_path.stat().st_size
            if file_size > self.max_file_size_bytes:
                return ExtractionResult(
                    success=False,
                    text="",
                    metadata={"file_size": file_size},
                    error_message=f"File too large: {file_size / (1024*1024):.1f}MB (max: {self.max_file_size_bytes / (1024*1024):.1f}MB)"
                )
            
            # Determine document type
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_types:
                return ExtractionResult(
                    success=False,
                    text="",
                    metadata={"file_extension": file_extension},
                    error_message=f"Unsupported file type: {file_extension}. Supported types: {list(self.supported_types.keys())}"
                )
            
            doc_type = self.supported_types[file_extension]
            
            # Read file content
            with open(file_path, 'rb') as file:
                file_content = file.read()
            
            # Extract text based on document type
            return self.extract_from_bytes(
                file_content=file_content,
                document_type=doc_type,
                filename=file_path.name
            )
            
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {str(e)}")
            return ExtractionResult(
                success=False,
                text="",
                metadata={},
                error_message=f"Extraction failed: {str(e)}"
            )
    
    def extract_from_bytes(
        self, 
        file_content: bytes, 
        document_type: DocumentType, 
        filename: str = "unknown"
    ) -> ExtractionResult:
        """
        Extract text from file content bytes.
        
        Args:
            file_content: Raw file content as bytes
            document_type: Type of document to process
            filename: Original filename for metadata
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        try:
            # Check file size
            if len(file_content) > self.max_file_size_bytes:
                return ExtractionResult(
                    success=False,
                    text="",
                    metadata={"file_size": len(file_content)},
                    error_message=f"File too large: {len(file_content) / (1024*1024):.1f}MB"
                )
            
            # Extract text based on document type
            if document_type == DocumentType.PDF:
                result = self._extract_pdf(file_content, filename)
            elif document_type == DocumentType.TXT:
                result = self._extract_txt(file_content, filename)
            elif document_type == DocumentType.DOCX:
                result = self._extract_docx(file_content, filename)
            elif document_type in [DocumentType.MD, DocumentType.MARKDOWN]:
                result = self._extract_markdown(file_content, filename)
            else:
                return ExtractionResult(
                    success=False,
                    text="",
                    metadata={},
                    error_message=f"Unsupported document type: {document_type}"
                )
            
            # Post-process extracted text
            if result.success and result.text:
                result.text = self._clean_text(result.text)
                # Recalculate statistics after cleaning
                result.word_count = len(result.text.split())
                result.character_count = len(result.text)
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {document_type.value} document: {str(e)}")
            return ExtractionResult(
                success=False,
                text="",
                metadata={"filename": filename, "document_type": document_type.value},
                error_message=f"Extraction failed: {str(e)}"
            )
    
    def _extract_pdf(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from PDF content."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                return ExtractionResult(
                    success=False,
                    text="",
                    metadata={"filename": filename, "encrypted": True},
                    error_message="PDF is encrypted and cannot be processed"
                )
            
            # Extract text from all pages
            text_parts = []
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1} of {filename}: {str(e)}")
                    continue
            
            extracted_text = "\n\n".join(text_parts)
            
            # Get PDF metadata
            metadata = {
                "filename": filename,
                "document_type": "pdf",
                "page_count": page_count,
                "encrypted": False
            }
            
            # Add PDF info if available
            if pdf_reader.metadata:
                pdf_info = pdf_reader.metadata
                metadata.update({
                    "title": pdf_info.get("/Title", ""),
                    "author": pdf_info.get("/Author", ""),
                    "subject": pdf_info.get("/Subject", ""),
                    "creator": pdf_info.get("/Creator", ""),
                    "producer": pdf_info.get("/Producer", ""),
                    "creation_date": str(pdf_info.get("/CreationDate", "")),
                    "modification_date": str(pdf_info.get("/ModDate", ""))
                })
            
            return ExtractionResult(
                success=True,
                text=extracted_text,
                metadata=metadata
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                metadata={"filename": filename, "document_type": "pdf"},
                error_message=f"PDF extraction failed: {str(e)}"
            )
    
    def _extract_txt(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from TXT content."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    
                    metadata = {
                        "filename": filename,
                        "document_type": "txt",
                        "encoding": encoding,
                        "line_count": len(text.splitlines())
                    }
                    
                    return ExtractionResult(
                        success=True,
                        text=text,
                        metadata=metadata
                    )
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail
            return ExtractionResult(
                success=False,
                text="",
                metadata={"filename": filename, "document_type": "txt"},
                error_message="Could not decode text file with any supported encoding"
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                metadata={"filename": filename, "document_type": "txt"},
                error_message=f"TXT extraction failed: {str(e)}"
            )
    
    def _extract_docx(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from DOCX content."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            # Extract text from paragraphs
            text_parts = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(text_parts)
            
            # Get document properties
            metadata = {
                "filename": filename,
                "document_type": "docx",
                "paragraph_count": paragraph_count,
                "table_count": table_count
            }
            
            # Add document core properties if available
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else ""
                })
            
            return ExtractionResult(
                success=True,
                text=extracted_text,
                metadata=metadata
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                metadata={"filename": filename, "document_type": "docx"},
                error_message=f"DOCX extraction failed: {str(e)}"
            )
    
    def _extract_markdown(self, file_content: bytes, filename: str) -> ExtractionResult:
        """Extract text from Markdown content."""
        try:
            # Decode markdown content
            encodings = ['utf-8', 'utf-16', 'latin-1']
            
            for encoding in encodings:
                try:
                    markdown_text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                return ExtractionResult(
                    success=False,
                    text="",
                    metadata={"filename": filename, "document_type": "markdown"},
                    error_message="Could not decode markdown file"
                )
            
            # Convert markdown to HTML then extract plain text
            html = markdown.markdown(markdown_text)
            
            # Simple HTML tag removal for plain text
            plain_text = re.sub(r'<[^>]+>', '', html)
            plain_text = re.sub(r'\n\s*\n', '\n\n', plain_text)  # Clean up extra newlines
            
            # Count markdown elements
            header_count = len(re.findall(r'^#+\s', markdown_text, re.MULTILINE))
            link_count = len(re.findall(r'\[([^\]]+)\]\([^)]+\)', markdown_text))
            code_block_count = len(re.findall(r'```', markdown_text)) // 2
            
            metadata = {
                "filename": filename,
                "document_type": "markdown",
                "header_count": header_count,
                "link_count": link_count,
                "code_block_count": code_block_count,
                "line_count": len(markdown_text.splitlines())
            }
            
            return ExtractionResult(
                success=True,
                text=plain_text,
                metadata=metadata
            )
            
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                metadata={"filename": filename, "document_type": "markdown"},
                error_message=f"Markdown extraction failed: {str(e)}"
            )
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines but preserve paragraph breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        return text
    
    def get_supported_extensions(self) -> list:
        """Get list of supported file extensions."""
        return list(self.supported_types.keys())
    
    def is_supported_file(self, filename: str) -> bool:
        """
        Check if a file is supported based on its extension.
        
        Args:
            filename: Name of the file to check
            
        Returns:
            True if file type is supported
        """
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.supported_types
    
    def get_document_type(self, filename: str) -> Optional[DocumentType]:
        """
        Get document type from filename.
        
        Args:
            filename: Name of the file
            
        Returns:
            DocumentType if supported, None otherwise
        """
        file_extension = Path(filename).suffix.lower()
        return self.supported_types.get(file_extension)